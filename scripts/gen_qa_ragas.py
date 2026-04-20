import argparse
import json
import os
import traceback
from pathlib import Path


def _approx_tokens(text: str) -> int:
    text = text or ""
    return max(1, len(text) // 2)


def load_docs_from_docx(docx_path, min_tokens=120):
    from docx import Document as DocxDocument
    from langchain_core.documents import Document

    path = Path(docx_path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX not found: {path}")

    doc = DocxDocument(str(path))

    blocks = []
    for p in doc.paragraphs:
        txt = (p.text or "").strip()
        if txt:
            blocks.append(txt)

    for table in doc.tables:
        row_texts = []
        for row in table.rows:
            cells = [" ".join((c.text or "").split()) for c in row.cells]
            if any(cells):
                row_texts.append(" | ".join(cells))
        if row_texts:
            blocks.append("\n".join(row_texts))

    docs = []
    buf = []
    sec = 1
    for b in blocks:
        buf.append(b)
        joined = "\n".join(buf)
        if _approx_tokens(joined) >= min_tokens:
            docs.append(
                Document(
                    page_content=joined,
                    metadata={"source": "docx", "source_file": path.name, "section": sec},
                )
            )
            sec += 1
            buf = []

    if buf:
        joined = "\n".join(buf)
        if _approx_tokens(joined) >= max(40, min_tokens // 2):
            docs.append(
                Document(
                    page_content=joined,
                    metadata={"source": "docx", "source_file": path.name, "section": sec},
                )
            )

    return docs


def _resolve_openai_config(openai_api_key_env, openai_base_url):
    api_key = os.getenv(openai_api_key_env)
    if not api_key:
        raise RuntimeError(f"{openai_api_key_env} is not set.")

    base_url = (
        openai_base_url
        or os.getenv("OPENAI_BASE_URL")
        or os.getenv("OPENAI_API_BASE")
        or None
    )
    return api_key, base_url


class _PlainOpenAIEmbeddings:
    def __init__(self, model, api_key, base_url=None):
        from openai import OpenAI

        self.model = model
        self.client = OpenAI(api_key=api_key, base_url=base_url)

    def embed_documents(self, texts):
        texts = [str(t) for t in texts]
        resp = self.client.embeddings.create(model=self.model, input=texts)
        return [x.embedding for x in resp.data]

    def embed_query(self, text):
        resp = self.client.embeddings.create(model=self.model, input=str(text))
        return resp.data[0].embedding

    async def aembed_documents(self, texts):
        return self.embed_documents(texts)

    async def aembed_query(self, text):
        return self.embed_query(text)


def _build_llm(provider, ollama_url, llm_model, openai_api_key_env, openai_base_url):
    from ragas.llms import LangchainLLMWrapper

    if provider == "openai":
        from langchain_openai import ChatOpenAI

        api_key, base_url = _resolve_openai_config(openai_api_key_env, openai_base_url)
        kwargs = {
            "model": llm_model,
            "temperature": 0,
            "api_key": api_key,
        }
        if base_url:
            kwargs["base_url"] = base_url

        return LangchainLLMWrapper(ChatOpenAI(**kwargs))

    from langchain_community.chat_models import ChatOllama

    return LangchainLLMWrapper(
        ChatOllama(
            model=llm_model,
            base_url=ollama_url,
            temperature=0,
        )
    )


def _build_emb(emb_provider, ollama_url, emb_model, openai_api_key_env, openai_base_url):
    from ragas.embeddings import LangchainEmbeddingsWrapper

    if emb_provider == "openai":
        api_key, base_url = _resolve_openai_config(openai_api_key_env, openai_base_url)
        plain = _PlainOpenAIEmbeddings(model=emb_model, api_key=api_key, base_url=base_url)
        return LangchainEmbeddingsWrapper(plain)

    from langchain_community.embeddings import OllamaEmbeddings

    return LangchainEmbeddingsWrapper(
        OllamaEmbeddings(
            model=emb_model,
            base_url=ollama_url,
        )
    )


def build_wrapped_models(provider, emb_provider, ollama_url, llm_model, emb_model, openai_api_key_env, openai_base_url):
    llm = _build_llm(provider, ollama_url, llm_model, openai_api_key_env, openai_base_url)
    emb = _build_emb(emb_provider, ollama_url, emb_model, openai_api_key_env, openai_base_url)
    return llm, emb


def generate_testset(docs, llm, emb, size):
    try:
        from ragas.testset import TestsetGenerator

        gen = TestsetGenerator(llm=llm, embedding_model=emb)
        testset = gen.generate_with_langchain_docs(
            documents=docs,
            testset_size=size,
            raise_exceptions=True,
        )
        return testset.to_pandas()
    except Exception as e:
        raise RuntimeError(
            "Failed in ragas.testset.TestsetGenerator.\n"
            f"{e}\n{traceback.format_exc()}"
        )


def normalize_rows(rows):
    out = []
    for i, row in enumerate(rows, 1):
        question = str(            (                row.get("question")                or row.get("query")                or row.get("user_input")                or row.get("input")                or row.get("prompt")                or ""            )        ).strip()
        answer = (            row.get("ground_truth")            or row.get("answer")            or row.get("reference")            or row.get("response")            or row.get("output")            or ""        )
        answer = str(answer).strip()

        keywords = []
        contexts = row.get("contexts") or row.get("retrieved_contexts") or []
        if isinstance(contexts, list) and contexts:
            snippet = str(contexts[0])[:120]
            keywords = [x for x in snippet.replace("，", " ").replace("。", " ").split(" ") if len(x) >= 2][:4]

        if not question:
            continue
        out.append(
            {
                "id": f"ragas_{i:04d}",
                "question": question,
                "gold_answer": answer,
                "evidence_keywords": keywords,
                "gold_page": None,
                "source": "ragas",
            }
        )
    return out


def main():
    parser = argparse.ArgumentParser(description="Generate QA candidates with Ragas from original DOCX.")
    parser.add_argument("--docx", default="精益六西格玛补充方案.docx")
    parser.add_argument("--out", default="data/qa_candidates_ragas.json")
    parser.add_argument("--size", type=int, default=120)
    parser.add_argument("--min-tokens", type=int, default=120)
    parser.add_argument("--provider", choices=["ollama", "openai"], default="ollama")
    parser.add_argument("--emb-provider", choices=["ollama", "openai"], default="ollama")
    parser.add_argument("--ollama-url", default="http://127.0.0.1:11434")
    parser.add_argument("--llm-model", default="qwen3:8b")
    parser.add_argument("--emb-model", default="bge-m3")
    parser.add_argument("--openai-api-key-env", default="OPENAI_API_KEY")
    parser.add_argument("--openai-base-url", default=None)
    args = parser.parse_args()

    docs = load_docs_from_docx(args.docx, min_tokens=args.min_tokens)
    if not docs:
        raise RuntimeError("No valid long docs loaded from original DOCX. Try lowering --min-tokens.")
    print(f"Loaded {len(docs)} long docs from {args.docx}.")

    llm, emb = build_wrapped_models(
        provider=args.provider,
        emb_provider=args.emb_provider,
        ollama_url=args.ollama_url,
        llm_model=args.llm_model,
        emb_model=args.emb_model,
        openai_api_key_env=args.openai_api_key_env,
        openai_base_url=args.openai_base_url,
    )
    df = generate_testset(docs, llm, emb, args.size)
    rows = df.to_dict(orient="records")
    raw_path = Path(args.out).with_name(Path(args.out).stem + "_raw.json")
    raw_path.parent.mkdir(parents=True, exist_ok=True)
    raw_path.write_text(json.dumps(rows, ensure_ascii=False, indent=2), encoding="utf-8")
    normalized = normalize_rows(rows)

    out_path = Path(args.out)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(json.dumps(normalized, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Generated {len(normalized)} ragas QA candidates -> {out_path}")


if __name__ == "__main__":
    main()


